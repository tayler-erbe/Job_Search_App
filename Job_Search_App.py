
import faiss
import numpy as np
import pandas as pd
import streamlit as st
from sentence_transformers import SentenceTransformer
import requests
import io
from docx import Document  # For Word document creation

# GitHub URLs for the required files
URL_CSV = "https://raw.githubusercontent.com/tayler-erbe/Job_Search_App/main/job_content_fixed.csv"
URL_EMBEDDINGS = "https://raw.githubusercontent.com/tayler-erbe/Job_Search_App/main/bert_embeddingsx.npy"
URL_FAISS_INDEX = "https://raw.githubusercontent.com/tayler-erbe/Job_Search_App/main/faiss_indexx.index"

# Load the job content data from GitHub
job_content_df = pd.read_csv(URL_CSV)

# Preprocess the 'Combine_String' column, filling NaN values with empty strings
combine_strings = job_content_df['Combine_String'].fillna('')

# Load the saved BERT embeddings from GitHub
response = requests.get(URL_EMBEDDINGS)
response.raise_for_status()
bert_embeddings = np.load(io.BytesIO(response.content))

# Load the saved FAISS index from GitHub
response = requests.get(URL_FAISS_INDEX)
response.raise_for_status()
with open("faiss_indexx.index", "wb") as f:
    f.write(response.content)

# Load FAISS index from the local file
faiss_index = faiss.read_index("faiss_indexx.index")

# Function to retrieve similar jobs using BERT embeddings
def retrieve_similar_jobs_bert(new_job_summary, faiss_index, bert_model, top_k=5):
    # Convert new job summary to BERT embedding
    new_embedding = bert_model.encode([new_job_summary], convert_to_tensor=True)
    
    # Use FAISS to find the top k most similar jobs
    distances, top_k_indices = faiss_index.search(new_embedding.cpu().numpy(), top_k)
    
    # Retrieve corresponding job descriptions
    similar_jobs = job_content_df.iloc[top_k_indices[0]].copy()
    similar_jobs['Similarity Score'] = 1 / (1 + distances[0])  # Convert distance to similarity score
    return similar_jobs

# Function to create a text file for download
def create_text_file(selected_job):
    output_text = f"""Job Title: {selected_job['Class Title']}

Job Duties: {selected_job['Job Duties']}

Minimum Qualifications: {selected_job['MAQ']}

Levels of Work: {selected_job['Levels Of Work']}

Knowledge, Skills, Abilities (KSA): {selected_job['KSA']}
"""
    return output_text

# Function to create a Word document for download
def create_word_doc(selected_job):
    doc = Document()
    doc.add_heading(selected_job['Class Title'], 0)
    doc.add_paragraph(f"Job Duties: {selected_job['Job Duties']}")
    doc.add_paragraph(f"Minimum Qualifications: {selected_job['MAQ']}")
    doc.add_paragraph(f"Levels of Work: {selected_job['Levels Of Work']}")
    doc.add_paragraph(f"Knowledge, Skills, Abilities (KSA): {selected_job['KSA']}")
    return doc

# Streamlit App
st.title("BERT and FAISS-Powered Job Search Tool")

# User Input for Job Query
user_query = st.text_input("Enter a job query (e.g., 'I need an accountant with 4 years of experience'): ")

# Allow user to select the number of returned job results (k)
top_k = st.number_input("Enter number of job results to return", min_value=1, max_value=20, value=5)

# Retrieve jobs based on the user query
if user_query:
    # Load a pre-trained BERT model for embedding new queries
    bert_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    
    # Display possible jobs retrieved using BERT and FAISS
    similar_jobs_bert = retrieve_similar_jobs_bert(user_query, faiss_index, bert_model, top_k=top_k)
    
    # Display a selection box for the user to pick a job title
    selected_job_title = st.selectbox("Select a job title", similar_jobs_bert['Class Title'].tolist())
    
    # Display full job posting when a job title is selected
    if selected_job_title:
        selected_job = similar_jobs_bert[similar_jobs_bert['Class Title'] == selected_job_title].iloc[0]
        
        # Display job posting details
        st.write(f"### {selected_job['Class Title']} (Similarity Score: {selected_job['Similarity Score']:.4f})")
        st.write(f"**Job Duties:** {selected_job['Job Duties']}")
        st.write(f"**Minimum Qualifications:** {selected_job['MAQ']}")
        st.write(f"**Levels of Work:** {selected_job['Levels Of Work']}")
        st.write(f"**Knowledge, Skills, Abilities (KSA):** {selected_job['KSA']}")
        
        # Create a downloadable text file
        text_output = create_text_file(selected_job)
        st.download_button(
            label="Download as Text File",
            data=text_output,
            file_name=f"{selected_job['Class Title']}.txt",
            mime="text/plain"
        )
        
        # Create a downloadable Word document
        word_doc = create_word_doc(selected_job)
        word_buffer = io.BytesIO()
        word_doc.save(word_buffer)
        word_buffer.seek(0)
        st.download_button(
            label="Download as Word Document",
            data=word_buffer,
            file_name=f"{selected_job['Class Title']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
